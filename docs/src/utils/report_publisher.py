"""
TradeMatrix.ai – ReportPublisher
--------------------------------
Generates and uploads reports to Google Drive and GitHub.
"""

from docx import Document
from datetime import datetime

class ReportPublisher:
    def __init__(self, output_path="data/weekly_briefing.docx"):
        self.output_path = output_path

    def create_report(self, trades):
        doc = Document()
        doc.add_heading("TradeMatrix.ai – Weekly Briefing", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        for t in trades:
            doc.add_paragraph(
                f"{t['asset']} | {t['direction']} | Entry {t['entry']} | TP {t['tp']} | SL {t['sl']} | Result: {t['result']}"
            )
        doc.save(self.output_path)
        print(f"Report saved → {self.output_path}")
