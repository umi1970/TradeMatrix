"""
TradeMatrix.ai - JournalBot Agent
Generates automated trading reports (PDF/DOCX) with AI-powered insights.

Execution: Daily at 21:00 MEZ (via Celery scheduler)
Report Types:
  1. Daily Report - Today's trades summary with AI analysis
  2. Weekly Report - Week's performance and trends
  3. Monthly Report - Comprehensive monthly review

Report Sections:
  - Trade Summary (win/loss ratio, P&L)
  - AI Insights (LangChain analysis of trading patterns)
  - Trade List (detailed table)
  - Performance Metrics (charts/stats)
  - Recommendations (AI-generated)

Output Formats: PDF (reportlab), DOCX (python-docx)
Storage: Supabase Storage (bucket: 'reports')
Database: Report metadata in 'reports' table
"""

import logging
from datetime import datetime, timedelta, timezone, date
from typing import Optional, Dict, Any, List
from decimal import Decimal
from uuid import UUID
import pytz
import os
import tempfile
import io

from supabase import Client
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import httpx

from openai import OpenAI

from src.chart_generator import ChartGenerator
from src.exceptions.chart_errors import (
    RateLimitError,
    ChartGenerationError,
    SymbolNotFoundError
)

# Setup logger
logger = logging.getLogger(__name__)


class JournalBot:
    """
    JournalBot Agent - Generates automated trading reports with AI insights

    Responsibilities:
    - Fetch trades data for specified period
    - Calculate performance metrics (win rate, P&L, R-multiples)
    - Generate AI summaries using LangChain + OpenAI
    - Create DOCX reports (python-docx)
    - Create PDF reports (reportlab)
    - Upload reports to Supabase Storage
    - Save report metadata to database
    """

    def __init__(self, supabase_client: Client, openai_api_key: str):
        """
        Initialize JournalBot agent

        Args:
            supabase_client: Supabase client instance (admin client for bypassing RLS)
            openai_api_key: OpenAI API key for GPT-4
        """
        self.supabase = supabase_client
        self.openai_api_key = openai_api_key
        self.chart_generator = ChartGenerator()

        # Initialize OpenAI client (direct API, no LangChain)
        self.openai_client = OpenAI(api_key=openai_api_key)

        logger.info("JournalBot initialized with ChartGenerator")


    def fetch_trades(
        self,
        user_id: Optional[UUID] = None,
        date_range: tuple = None
    ) -> List[Dict[str, Any]]:
        """
        Fetch trades for specified period

        Args:
            user_id: Optional user ID (None = all users)
            date_range: Tuple of (start_date, end_date) as datetime objects
                       Default: today

        Returns:
            List of trade records with related data
        """
        logger.info(f"Fetching trades for user_id={user_id}, date_range={date_range}")

        try:
            # Default to today if no date range provided
            if not date_range:
                berlin_tz = pytz.timezone('Europe/Berlin')
                today = datetime.now(berlin_tz).date()
                date_range = (
                    datetime.combine(today, datetime.min.time()),
                    datetime.combine(today, datetime.max.time())
                )

            start_date, end_date = date_range

            # Build query
            query = self.supabase.table('trades')\
                .select('*, market_symbols(symbol, alias)')\
                .gte('created_at', start_date.isoformat())\
                .lte('created_at', end_date.isoformat())\
                .order('created_at', desc=True)

            # Filter by user if specified
            if user_id:
                query = query.eq('user_id', str(user_id))

            result = query.execute()

            trades = result.data if result.data else []
            logger.info(f"Fetched {len(trades)} trades")

            return trades

        except Exception as e:
            logger.error(f"Error fetching trades: {str(e)}", exc_info=True)
            return []


    def generate_ai_summary(self, trades_data: List[Dict[str, Any]]) -> Dict[str, str]:
        """
        Generate AI-powered summary and insights using OpenAI GPT-4

        Args:
            trades_data: List of trade records

        Returns:
            Dictionary with AI-generated content:
            {
                'summary': 'Overall trading performance summary',
                'insights': 'Key patterns and observations',
                'recommendations': 'Actionable suggestions'
            }
        """
        logger.info(f"Generating AI summary for {len(trades_data)} trades")

        try:
            # Calculate basic metrics for context
            total_trades = len(trades_data)
            winning_trades = [t for t in trades_data if t.get('pnl', 0) > 0]
            losing_trades = [t for t in trades_data if t.get('pnl', 0) < 0]

            win_rate = (len(winning_trades) / total_trades * 100) if total_trades > 0 else 0
            total_pnl = sum(float(t.get('pnl', 0)) for t in trades_data)

            avg_win = sum(float(t.get('pnl', 0)) for t in winning_trades) / len(winning_trades) if winning_trades else 0
            avg_loss = sum(float(t.get('pnl', 0)) for t in losing_trades) / len(losing_trades) if losing_trades else 0

            # Prepare context for LLM
            trades_summary = "\n".join([
                f"- {t.get('side', 'N/A').upper()} {t.get('market_symbols', {}).get('symbol', 'N/A')} @ {t.get('entry_price', 0)}, "
                f"Exit: {t.get('exit_price', 'N/A')}, P&L: {t.get('pnl', 0)}, "
                f"Strategy: {t.get('strategy', 'N/A')}"
                for t in trades_data[:20]  # Limit to 20 trades for context window
            ])

            # Generate summary using OpenAI API (direct, no LangChain)
            prompt_text = f"""Analyze these trading statistics and trades:

Total Trades: {total_trades}
Win Rate: {win_rate:.1f}%
Total P&L: {total_pnl:.2f}
Average Win: {avg_win:.2f}
Average Loss: {avg_loss:.2f}

Recent Trades:
{trades_summary}

Provide:
1. SUMMARY: 2-3 sentence overview of performance
2. INSIGHTS: Key patterns or observations (3-4 bullet points)
3. RECOMMENDATIONS: Actionable suggestions for improvement (3-4 bullet points)

Format as JSON with keys: summary, insights, recommendations"""

            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                temperature=0.7,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert trading analyst for TradeMatrix.ai. Analyze the trading performance data and provide professional insights. Focus on: patterns, execution quality, risk management, and actionable improvements."
                    },
                    {
                        "role": "user",
                        "content": prompt_text
                    }
                ]
            )

            # Parse response
            content = response.choices[0].message.content

            # Try to extract structured data (fallback to simple parsing)
            try:
                import json
                parsed = json.loads(content)
                ai_summary = {
                    'summary': parsed.get('summary', f"Total Trades: {total_trades} | Win Rate: {win_rate:.1f}% | P&L: {total_pnl:.2f}"),
                    'insights': parsed.get('insights', content),
                    'recommendations': parsed.get('recommendations', "Continue following trading rules and risk management principles.")
                }
            except json.JSONDecodeError:
                # Fallback if not valid JSON
                ai_summary = {
                    'summary': f"Total Trades: {total_trades} | Win Rate: {win_rate:.1f}% | P&L: {total_pnl:.2f}",
                    'insights': content,
                    'recommendations': "Continue following trading rules and risk management principles."
                }

            logger.info("AI summary generated successfully")
            return ai_summary

        except Exception as e:
            logger.error(f"Error generating AI summary: {str(e)}", exc_info=True)

            # Return fallback summary
            return {
                'summary': f"Analyzed {len(trades_data)} trades",
                'insights': "Unable to generate AI insights at this time.",
                'recommendations': "Continue monitoring performance and following trading rules."
            }


    def create_docx_report(
        self,
        report_data: Dict[str, Any],
        filename: str
    ) -> Optional[str]:
        """
        Generate DOCX report using python-docx

        Args:
            report_data: Dictionary with report content
                {
                    'title': 'Daily Trading Report',
                    'date': datetime,
                    'metrics': {...},
                    'ai_summary': {...},
                    'trades': [...]
                }
            filename: Output filename (without extension)

        Returns:
            Path to created DOCX file or None on error
        """
        logger.info(f"Creating DOCX report: {filename}")

        try:
            # Create document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            # Title
            title = doc.add_heading(report_data.get('title', 'Trading Report'), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Date
            date_str = report_data.get('date', datetime.now()).strftime('%Y-%m-%d')
            date_para = doc.add_paragraph(f"Report Date: {date_str}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacer

            # Performance Metrics Section
            doc.add_heading('Performance Summary', level=2)

            metrics = report_data.get('metrics', {})
            metrics_table = doc.add_table(rows=6, cols=2)
            metrics_table.style = 'Light Grid Accent 1'

            metrics_data = [
                ('Total Trades', str(metrics.get('total_trades', 0))),
                ('Winning Trades', str(metrics.get('winning_trades', 0))),
                ('Losing Trades', str(metrics.get('losing_trades', 0))),
                ('Win Rate', f"{metrics.get('win_rate', 0):.1f}%"),
                ('Total P&L', f"{metrics.get('total_pnl', 0):.2f}"),
                ('Average R-Multiple', f"{metrics.get('avg_r_multiple', 0):.2f}R"),
            ]

            for idx, (label, value) in enumerate(metrics_data):
                metrics_table.rows[idx].cells[0].text = label
                metrics_table.rows[idx].cells[1].text = value

            doc.add_paragraph()  # Spacer

            # EOD Performance Section (if available)
            eod_performance = metrics.get('eod_performance', {})
            if eod_performance.get('status') == 'success' and eod_performance.get('symbols'):
                doc.add_heading('Yesterday Market Performance', level=2)

                eod_table = doc.add_table(rows=1, cols=5)
                eod_table.style = 'Light Grid Accent 1'

                # Header row
                header_cells = eod_table.rows[0].cells
                headers = ['Symbol', 'Y. High', 'Y. Low', 'Y. Close', 'Change %']
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    header_cells[idx].paragraphs[0].runs[0].font.bold = True

                # Data rows
                for symbol in eod_performance['symbols']:
                    row_cells = eod_table.add_row().cells
                    row_cells[0].text = symbol['symbol']
                    row_cells[1].text = f"{symbol['yesterday_high']:.2f}"
                    row_cells[2].text = f"{symbol['yesterday_low']:.2f}"
                    row_cells[3].text = f"{symbol['yesterday_close']:.2f}"

                    change_pct = symbol['daily_change_percent']
                    row_cells[4].text = f"{change_pct:+.2f}%"

                    # Color code change
                    if change_pct > 0:
                        row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif change_pct < 0:
                        row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red

                doc.add_paragraph()  # Spacer

            # AI Insights Section
            doc.add_heading('AI-Powered Insights', level=2)

            ai_summary = report_data.get('ai_summary', {})

            doc.add_heading('Summary', level=3)
            doc.add_paragraph(ai_summary.get('summary', 'No summary available'))

            doc.add_heading('Key Insights', level=3)
            doc.add_paragraph(ai_summary.get('insights', 'No insights available'))

            doc.add_heading('Recommendations', level=3)
            doc.add_paragraph(ai_summary.get('recommendations', 'No recommendations available'))

            doc.add_paragraph()  # Spacer

            # Trades List Section
            doc.add_heading('Trade Details', level=2)

            trades = report_data.get('trades', [])

            if trades:
                trades_table = doc.add_table(rows=1, cols=6)
                trades_table.style = 'Light Grid Accent 1'

                # Header row
                header_cells = trades_table.rows[0].cells
                headers = ['Time', 'Symbol', 'Side', 'Entry', 'Exit', 'P&L']
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    header_cells[idx].paragraphs[0].runs[0].font.bold = True

                # Data rows
                for trade in trades[:20]:  # Limit to 20 trades
                    row_cells = trades_table.add_row().cells

                    created_at = trade.get('created_at', '')
                    if created_at:
                        created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00')).strftime('%H:%M')

                    row_cells[0].text = created_at
                    row_cells[1].text = trade.get('market_symbols', {}).get('symbol', 'N/A')
                    row_cells[2].text = trade.get('side', 'N/A').upper()
                    row_cells[3].text = f"{trade.get('entry_price', 0):.2f}"
                    row_cells[4].text = f"{trade.get('exit_price', 0):.2f}" if trade.get('exit_price') else 'Open'

                    pnl = float(trade.get('pnl', 0))
                    row_cells[5].text = f"{pnl:+.2f}"

                    # Color code P&L
                    if pnl > 0:
                        row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif pnl < 0:
                        row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
            else:
                doc.add_paragraph("No trades for this period.")

            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph("Generated by TradeMatrix.ai JournalBot")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Save document
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"{filename}.docx")
            doc.save(file_path)

            logger.info(f"DOCX report created: {file_path}")
            return file_path

        except Exception as e:
            logger.error(f"Error creating DOCX report: {str(e)}", exc_info=True)
            return None


    def create_pdf_report(
        self,
        report_data: Dict[str, Any],
        filename: str
    ) -> Optional[str]:
        """
        Generate PDF report using reportlab

        Args:
            report_data: Dictionary with report content
            filename: Output filename (without extension)

        Returns:
            Path to created PDF file or None on error
        """
        logger.info(f"Creating PDF report: {filename}")

        try:
            # Setup PDF
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"{filename}.pdf")

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12,
                spaceBefore=12
            )

            # Title
            title = Paragraph(report_data.get('title', 'Trading Report'), title_style)
            story.append(title)

            # Date
            date_str = report_data.get('date', datetime.now()).strftime('%Y-%m-%d')
            date_para = Paragraph(f"Report Date: {date_str}", styles['Normal'])
            story.append(date_para)
            story.append(Spacer(1, 0.3*inch))

            # Performance Metrics
            story.append(Paragraph('Performance Summary', heading_style))

            metrics = report_data.get('metrics', {})
            metrics_data = [
                ['Metric', 'Value'],
                ['Total Trades', str(metrics.get('total_trades', 0))],
                ['Winning Trades', str(metrics.get('winning_trades', 0))],
                ['Losing Trades', str(metrics.get('losing_trades', 0))],
                [f'Win Rate', f"{metrics.get('win_rate', 0):.1f}%"],
                ['Total P&L', f"{metrics.get('total_pnl', 0):.2f}"],
                ['Avg R-Multiple', f"{metrics.get('avg_r_multiple', 0):.2f}R"],
            ]

            metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(metrics_table)
            story.append(Spacer(1, 0.3*inch))

            # AI Insights
            story.append(Paragraph('AI-Powered Insights', heading_style))

            ai_summary = report_data.get('ai_summary', {})

            story.append(Paragraph('<b>Summary</b>', styles['Normal']))
            story.append(Paragraph(ai_summary.get('summary', 'No summary available'), styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

            story.append(Paragraph('<b>Key Insights</b>', styles['Normal']))
            story.append(Paragraph(ai_summary.get('insights', 'No insights available'), styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

            story.append(Paragraph('<b>Recommendations</b>', styles['Normal']))
            story.append(Paragraph(ai_summary.get('recommendations', 'No recommendations available'), styles['Normal']))
            story.append(Spacer(1, 0.3*inch))

            # Trades List
            story.append(Paragraph('Trade Details', heading_style))

            trades = report_data.get('trades', [])
            trade_charts = report_data.get('trade_charts', [])

            if trades:
                trades_data = [['Time', 'Symbol', 'Side', 'Entry', 'Exit', 'P&L']]

                for trade in trades[:20]:  # Limit to 20 trades
                    created_at = trade.get('created_at', '')
                    if created_at:
                        created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00')).strftime('%H:%M')

                    pnl = float(trade.get('pnl', 0))

                    trades_data.append([
                        created_at,
                        trade.get('market_symbols', {}).get('symbol', 'N/A'),
                        trade.get('side', 'N/A').upper(),
                        f"{trade.get('entry_price', 0):.2f}",
                        f"{trade.get('exit_price', 0):.2f}" if trade.get('exit_price') else 'Open',
                        f"{pnl:+.2f}"
                    ])

                trades_table = Table(trades_data, colWidths=[0.8*inch, 1*inch, 0.8*inch, 1*inch, 1*inch, 1*inch])
                trades_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                story.append(trades_table)

                # Add charts section if available
                if trade_charts:
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph('Trade Charts', heading_style))

                    for chart_info in trade_charts[:5]:  # Limit to 5 charts in PDF
                        try:
                            # Download chart image
                            img_data = self._download_chart_image(chart_info['chart_url'])
                            if img_data:
                                # Create image from bytes
                                img = Image(io.BytesIO(img_data), width=5*inch, height=3*inch)

                                # Add chart label
                                chart_label = f"{chart_info['symbol']} - Entry: {chart_info.get('entry_price', 'N/A')}, P&L: {chart_info.get('pnl', 0):+.2f}"
                                story.append(Paragraph(chart_label, styles['Normal']))
                                story.append(img)
                                story.append(Spacer(1, 0.2*inch))
                        except Exception as e:
                            logger.error(f"Error adding chart to PDF: {e}")
                            continue

            else:
                story.append(Paragraph("No trades for this period.", styles['Normal']))

            # Footer
            story.append(Spacer(1, 0.3*inch))
            footer = Paragraph("Generated by TradeMatrix.ai JournalBot", styles['Normal'])
            story.append(footer)

            # Build PDF
            doc.build(story)

            logger.info(f"PDF report created: {file_path}")
            return file_path

        except Exception as e:
            logger.error(f"Error creating PDF report: {str(e)}", exc_info=True)
            return None


    def upload_to_storage(
        self,
        file_path: str,
        bucket: str = 'reports'
    ) -> Optional[str]:
        """
        Upload report file to Supabase Storage

        Args:
            file_path: Local path to file
            bucket: Supabase storage bucket name

        Returns:
            Public URL of uploaded file or None on error
        """
        logger.info(f"Uploading {file_path} to bucket: {bucket}")

        try:
            # Read file
            with open(file_path, 'rb') as f:
                file_data = f.read()

            # Generate storage path
            filename = os.path.basename(file_path)
            storage_path = f"reports/{datetime.now().strftime('%Y/%m')}/{filename}"

            # Upload to Supabase Storage
            result = self.supabase.storage.from_(bucket).upload(
                storage_path,
                file_data,
                file_options={"content-type": "application/octet-stream"}
            )

            # Get public URL
            public_url = self.supabase.storage.from_(bucket).get_public_url(storage_path)

            logger.info(f"File uploaded successfully: {public_url}")
            return public_url

        except Exception as e:
            logger.error(f"Error uploading file to storage: {str(e)}", exc_info=True)
            return None


    def create_report_record(self, report_metadata: Dict[str, Any]) -> Optional[UUID]:
        """
        Save report metadata to 'reports' table

        Args:
            report_metadata: Dictionary with report details
                {
                    'user_id': UUID or None,
                    'report_type': 'daily' | 'weekly' | 'monthly',
                    'report_date': date,
                    'title': str,
                    'content': str,
                    'ai_summary': str,
                    'ai_insights': dict or list,
                    'pdf_url': str,
                    'file_url_docx': str,
                    'metrics': {...}
                }

        Returns:
            UUID of created report record or None on error
        """
        logger.info(f"Creating report record for {report_metadata.get('report_type')}")

        try:
            # Prepare record with correct column names matching schema
            record = {
                'user_id': str(report_metadata['user_id']) if report_metadata.get('user_id') else None,
                'title': report_metadata.get('title', f"{report_metadata.get('report_type', 'daily').capitalize()} Report"),
                'content': report_metadata.get('content', ''),
                'report_type': report_metadata.get('report_type', 'daily'),
                'ai_summary': report_metadata.get('ai_summary'),
                'ai_insights': report_metadata.get('ai_insights', []),
                'status': report_metadata.get('status', 'completed'),  # Use 'completed' instead of 'draft' for generated reports
                'pdf_url': report_metadata.get('pdf_url'),
                'metrics': report_metadata.get('metrics', {}),  # Separate metrics column (Migration 015)
                'metadata': {
                    'report_date': report_metadata.get('report_date', datetime.now().date()).isoformat(),
                    'file_url_docx': report_metadata.get('file_url_docx'),  # DOCX URL in metadata
                    **report_metadata.get('metadata', {})
                },
                'created_at': datetime.now(timezone.utc).isoformat()
            }

            # Insert into database
            result = self.supabase.table('reports').insert(record).execute()

            if result.data and len(result.data) > 0:
                report_id = result.data[0]['id']
                logger.info(f"Report record created: {report_id}")
                return UUID(report_id)
            else:
                logger.error("Failed to create report record")
                return None

        except Exception as e:
            logger.error(f"Error creating report record: {str(e)}", exc_info=True)
            return None


    def _download_chart_image(self, chart_url: str) -> Optional[bytes]:
        """
        Download chart image from URL

        Args:
            chart_url: URL of the chart image

        Returns:
            Image bytes or None on error
        """
        try:
            with httpx.Client(timeout=30.0) as client:
                response = client.get(chart_url)
                response.raise_for_status()
                return response.content
        except Exception as e:
            logger.error(f"Error downloading chart image: {e}")
            return None

    def _generate_trade_charts(self, trades: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Generate charts for all trades

        Args:
            trades: List of trade records

        Returns:
            List of dicts with trade_id, symbol, chart_url, entry_price
        """
        logger.info(f"Generating charts for {len(trades)} trades...")

        report_charts = []

        for trade in trades[:20]:  # Limit to 20 trades for reports
            try:
                symbol_id = trade.get('symbol_id')
                if not symbol_id:
                    logger.warning(f"Trade {trade.get('id')} missing symbol_id - skipping chart")
                    continue

                # Generate chart for this trade
                snapshot = self.chart_generator.generate_chart(
                    symbol_id=symbol_id,
                    timeframe='4h',  # 4h for reports
                    trigger_type='report'
                )

                report_charts.append({
                    'trade_id': trade['id'],
                    'symbol': trade.get('market_symbols', {}).get('symbol', 'N/A'),
                    'chart_url': snapshot['chart_url'],
                    'chart_snapshot_id': snapshot.get('snapshot_id'),
                    'entry_price': trade.get('entry_price'),
                    'exit_price': trade.get('exit_price'),
                    'pnl': trade.get('pnl', 0)
                })

                logger.info(f"  ✅ Chart generated for {trade['id']}")

            except RateLimitError as e:
                logger.warning(f"Rate limit reached during chart generation: {e.details}")
                # Stop generating more charts
                break
            except (ChartGenerationError, SymbolNotFoundError) as e:
                logger.warning(f"Chart generation failed for trade {trade.get('id')}: {e}")
                # Continue without chart
                continue
            except Exception as e:
                logger.error(f"Unexpected error generating chart for trade {trade.get('id')}: {e}")
                continue

        logger.info(f"Generated {len(report_charts)} charts for report")
        return report_charts

    def _get_eod_performance(self, trade_date: date) -> Dict[str, Any]:
        """
        Fetch yesterday's market performance from EOD data

        Args:
            trade_date: Current trading date

        Returns:
            Dict with EOD performance metrics for all active symbols
        """
        try:
            # Query eod_data for yesterday (trade_date - 1 day)
            yesterday = trade_date - timedelta(days=1)
            trade_date_str = trade_date.strftime('%Y-%m-%d')

            response = self.supabase.table('eod_levels')\
                .select('*, symbols(symbol, name)')\
                .eq('trade_date', trade_date_str)\
                .execute()

            if not response.data:
                return {
                    'status': 'no_data',
                    'symbols': []
                }

            # Format EOD performance summary
            symbols_performance = []
            for level in response.data:
                if level.get('symbols'):
                    symbol_data = {
                        'symbol': level['symbols']['symbol'],
                        'name': level['symbols']['name'],
                        'yesterday_high': float(level.get('yesterday_high', 0)),
                        'yesterday_low': float(level.get('yesterday_low', 0)),
                        'yesterday_close': float(level.get('yesterday_close', 0)),
                        'daily_change_points': float(level.get('daily_change_points', 0)),
                        'daily_change_percent': float(level.get('daily_change_percent', 0))
                    }
                    symbols_performance.append(symbol_data)

            return {
                'status': 'success',
                'date': trade_date_str,
                'symbols': symbols_performance
            }

        except Exception as e:
            logger.error(f"Error fetching EOD performance: {str(e)}", exc_info=True)
            return {
                'status': 'error',
                'error': str(e),
                'symbols': []
            }

    def generate_daily_report(self, user_id: Optional[UUID] = None) -> Dict[str, Any]:
        """
        Generate daily trading report

        Main workflow:
        1. Fetch today's trades
        2. Calculate performance metrics
        3. Generate AI summary
        4. Create PDF and DOCX reports
        5. Upload to Supabase Storage
        6. Save report metadata to database

        Args:
            user_id: Optional user ID (None = all users)

        Returns:
            Dictionary with report generation results
        """
        logger.info(f"Generating daily report for user_id={user_id}")

        try:
            # Step 1 - Fetch today's trades
            berlin_tz = pytz.timezone('Europe/Berlin')
            today = datetime.now(berlin_tz).date()

            date_range = (
                datetime.combine(today, datetime.min.time(), tzinfo=berlin_tz),
                datetime.combine(today, datetime.max.time(), tzinfo=berlin_tz)
            )

            trades = self.fetch_trades(user_id=user_id, date_range=date_range)

            # Step 2 - Calculate metrics
            total_trades = len(trades)
            winning_trades = [t for t in trades if float(t.get('pnl', 0)) > 0]
            losing_trades = [t for t in trades if float(t.get('pnl', 0)) < 0]

            win_rate = (len(winning_trades) / total_trades * 100) if total_trades > 0 else 0
            total_pnl = sum(float(t.get('pnl', 0)) for t in trades)

            # Calculate R-multiples
            r_multiples = [float(t.get('r_multiple', 0)) for t in trades if t.get('r_multiple')]
            avg_r_multiple = sum(r_multiples) / len(r_multiples) if r_multiples else 0

            # Step 2b - Fetch EOD performance data (yesterday's market moves)
            eod_performance = self._get_eod_performance(today)

            metrics = {
                'total_trades': total_trades,
                'winning_trades': len(winning_trades),
                'losing_trades': len(losing_trades),
                'win_rate': round(win_rate, 1),
                'total_pnl': round(total_pnl, 2),
                'avg_r_multiple': round(avg_r_multiple, 2),
                'eod_performance': eod_performance  # Market context
            }

            # Step 3 - Generate charts for trades
            trade_charts = self._generate_trade_charts(trades)

            # Step 4 - Generate AI summary
            ai_summary = self.generate_ai_summary(trades)

            # Step 5 - Prepare report data
            report_data = {
                'title': f'Daily Trading Report - {today.strftime("%Y-%m-%d")}',
                'date': today,
                'metrics': metrics,
                'ai_summary': ai_summary,
                'trades': trades,
                'trade_charts': trade_charts  # Add charts to report data
            }

            # Step 6 - Create reports
            filename_base = f"daily_report_{today.strftime('%Y%m%d')}"
            if user_id:
                filename_base += f"_user_{str(user_id)[:8]}"

            pdf_path = self.create_pdf_report(report_data, filename_base)
            docx_path = self.create_docx_report(report_data, filename_base)

            # Step 7 - Upload to storage
            pdf_url = self.upload_to_storage(pdf_path) if pdf_path else None
            docx_url = self.upload_to_storage(docx_path) if docx_path else None

            # Step 8 - Save report metadata
            report_metadata = {
                'user_id': user_id,
                'title': report_data['title'],
                'content': f"Daily trading report for {today.strftime('%Y-%m-%d')} with {total_trades} trades analyzed.",
                'report_type': 'daily',
                'report_date': today,
                'ai_summary': ai_summary.get('summary', ''),
                'ai_insights': {
                    'insights': ai_summary.get('insights', ''),
                    'recommendations': ai_summary.get('recommendations', '')
                },
                'status': 'published',
                'pdf_url': pdf_url,  # Correct column name
                'file_url_docx': docx_url,
                'metrics': metrics  # Will be stored in metadata JSONB column
            }

            report_id = self.create_report_record(report_metadata)

            # Cleanup temp files
            if pdf_path and os.path.exists(pdf_path):
                os.remove(pdf_path)
            if docx_path and os.path.exists(docx_path):
                os.remove(docx_path)

            result = {
                'success': True,
                'report_id': str(report_id) if report_id else None,
                'report_date': today.isoformat(),
                'metrics': metrics,
                'pdf_url': pdf_url,
                'docx_url': docx_url,
                'trades_analyzed': total_trades
            }

            logger.info(f"Daily report generated successfully: {result}")
            return result

        except Exception as e:
            logger.error(f"Error generating daily report: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': str(e)
            }


    def run(self) -> Dict[str, Any]:
        """
        Main execution method - Called by Celery scheduler at 21:00 MEZ daily

        Process:
        1. Fetch all active users (or generate global report)
        2. For each user, generate daily report
        3. Return summary of generated reports

        Returns:
            Dictionary with execution summary:
            {
                'execution_time': datetime,
                'reports_generated': int,
                'reports': List[Dict]
            }
        """
        execution_start = datetime.now(timezone.utc)
        logger.info(f"JournalBot execution started at {execution_start}")

        try:
            # For now, generate a single global report (all users combined)
            # In production, you would fetch users and generate per-user reports

            report_result = self.generate_daily_report(user_id=None)

            reports = []
            if report_result.get('success'):
                reports.append(report_result)

            execution_end = datetime.now(timezone.utc)
            duration_ms = int((execution_end - execution_start).total_seconds() * 1000)

            summary = {
                'execution_time': execution_start.isoformat(),
                'execution_duration_ms': duration_ms,
                'reports_generated': len(reports),
                'reports': reports
            }

            logger.info(f"JournalBot execution completed: {summary}")
            return summary

        except Exception as e:
            logger.error(f"Fatal error in JournalBot execution: {str(e)}", exc_info=True)
            execution_end = datetime.now(timezone.utc)
            duration_ms = int((execution_end - execution_start).total_seconds() * 1000)

            return {
                'execution_time': execution_start.isoformat(),
                'execution_duration_ms': duration_ms,
                'reports_generated': 0,
                'reports': [],
                'error': str(e)
            }


# Example usage (for testing)
if __name__ == "__main__":
    # This would be called from Celery task
    import sys
    import os
    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../api/src')))

    from config.supabase import get_supabase_admin
    from config import settings

    # Initialize agent with admin client
    bot = JournalBot(
        supabase_client=get_supabase_admin(),
        openai_api_key=settings.OPENAI_API_KEY
    )

    # Run report generation
    result = bot.run()

    print("JournalBot Results:")
    print(result)
